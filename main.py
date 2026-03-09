import argparse
import json
import os
import re
from pathlib import Path
from typing import Any, Dict, List

import pdfplumber
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from docx2pdf import convert as docx2pdf_convert
except ImportError:  # graceful degradation if docx2pdf is missing
    docx2pdf_convert = None


def extract_text_from_pdf(pdf_path: Path) -> str:
    if not pdf_path.is_file():
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    texts: List[str] = []
    with pdfplumber.open(str(pdf_path)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            texts.append(page_text)
    return "\n\n".join(texts).strip()


def load_profile(profile_path: Path) -> Dict[str, Any]:
    if not profile_path.is_file():
        raise FileNotFoundError(f"Profile JSON not found: {profile_path}")
    with profile_path.open("r", encoding="utf-8") as f:
        return json.load(f)


# Professional, ATS-friendly styling
FONT_NAME = "Calibri"
BODY_PT = 11
HEADING_PT = 12
NAME_PT = 22
CONTACT_PT = 10
LINE_SPACING = 1.15  # Slightly open for readability
HEADER_SPACE_BELOW = Pt(14)
SECTION_HEADING_SPACE_ABOVE = Pt(14)
SECTION_HEADING_SPACE_BELOW = Pt(6)


def _set_run_font(run, size_pt: int = BODY_PT, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.bold = bold


def _add_horizontal_rule(document: Document, space_after: int = 12) -> None:
    """Add a thin gray line to separate header from body."""
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4472C4")  # Professional blue-gray
    pBdr.append(bottom)
    pPr.append(pBdr)


def _ats_section_heading(document: Document, text: str) -> None:
    """Section heading: bold, uppercase, with underline and consistent spacing."""
    p = document.add_paragraph()
    run = p.add_run(text.upper())
    _set_run_font(run, size_pt=HEADING_PT, bold=True)
    p.paragraph_format.space_before = SECTION_HEADING_SPACE_ABOVE
    p.paragraph_format.space_after = SECTION_HEADING_SPACE_BELOW
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # Underline for a clean, professional look
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2F5496")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _ats_body_paragraph(document: Document, text: str, space_after: int = 4) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size_pt=BODY_PT, bold=False)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def _ats_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    _set_run_font(run, size_pt=BODY_PT, bold=False)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


# Boilerplate: lines that are only these phrases (or with colons) are removed from final output
_BOILERPLATE_LINE = re.compile(
    r"^[:\s]*(resume|curriculum\s*vitae|cv|résumé|career\s*objective)[:\s]*$",
    re.IGNORECASE,
)
# Section headers / literals to drop (standalone lines)
_SECTION_LITERAL_LINE = re.compile(
    r"^[\s\uFFFD\u00A0]*(experience|career\s*objective|personal\s*details|"
    r"personal\s*information|education|skills|summary|work\s*experience|"
    r"professional\s*summary|qualification|references?)[\s:\uFFFD\u00A0]*$",
    re.IGNORECASE,
)
# Strip leading "Resume:", ":Resume:", etc. from any line
_LEADING_BOILERPLATE = re.compile(
    r"^[:\s]*(resume|curriculum\s*vitae|cv|résumé)[:\s]*",
    re.IGNORECASE,
)
# Line looks like a list item (bullet, dash, number)
_BULLET_LINE = re.compile(
    r"^[\s\uFFFD\u00A0]*([\u2022\u2023\u2043\u2219\u00B7\*\-\–\—]\s*|\d+[\.\)]\s*)",
    re.UNICODE,
)
# Labels that appear in header — lines containing only these (one or more "Label : value") are dropped to avoid duplication
_HEADER_LABELS = re.compile(
    r"(Name|Contact\s*no|Email\s*id|Date\s*of\s*birth|DOB|Address|Gender)\s*:\s*",
    re.IGNORECASE,
)


def _strip_bad_chars(s: str) -> str:
    """Remove replacement char, control chars, and other non-printable from string."""
    if not s:
        return ""
    # Replace U+FFFD (replacement character) and control chars with space
    s = "".join(c if c.isprintable() or c in "\n\t" else " " for c in s)
    s = s.replace("\uFFFD", " ").replace("\u00A0", " ")
    return re.sub(r"\s+", " ", s).strip()


def _is_duplicate_header_line(line: str) -> bool:
    """True if line is only header-style fields (Name, Contact no, Email, DOB, Address, Gender) — already at top of CV."""
    s = line.strip()
    if not s:
        return True
    # Single label:value for a header field
    if re.match(r"^(Name|Contact\s*no|Email\s*id|Date\s*of\s*birth|DOB|Address|Gender)\s*:\s*.+$", s, re.IGNORECASE):
        return True
    # Line is only a sequence of "Label : value" for header fields (e.g. "Contact no : 9624635666 Email id : a@b.com")
    remainder = s
    while remainder:
        m = _HEADER_LABELS.match(remainder)
        if not m:
            return False
        rest = remainder[m.end() :].strip()
        next_m = _HEADER_LABELS.search(rest)
        if next_m:
            remainder = rest[next_m.start() :].strip()
        else:
            remainder = ""
    return True


def _clean_extracted_cv_text(text: str, name_at_top: str = "") -> str:
    """Remove boilerplate, literals, duplicate header fields, and bad Unicode from extracted CV text."""
    if not text or not text.strip():
        return ""
    name_upper = (name_at_top or "").strip().upper()
    lines = []
    for raw_line in text.splitlines():
        line = _strip_bad_chars(raw_line)
        if not line:
            continue
        if _BOILERPLATE_LINE.match(line):
            continue
        if _SECTION_LITERAL_LINE.match(line):
            continue
        line = _LEADING_BOILERPLATE.sub("", line).strip()
        line = _strip_bad_chars(line)
        if not line:
            continue
        # Drop lines that are only "Name :", "Email :" etc. (label-only)
        if re.match(r"^[A-Za-z\s]+:\s*$", line):
            continue
        # Drop short all-caps fragments (e.g. "CARE", "Perso" from broken headers)
        if len(line) <= 25 and line.isupper():
            continue
        if len(line) <= 20 and re.match(r"^[A-Z][a-z]*\s*$", line):
            continue
        # Drop lines that duplicate header (Name, Contact no, Email, DOB, Address, Gender — already at top)
        if _is_duplicate_header_line(line):
            continue
        # Drop "Name + Address" line (e.g. "MAKWANA MAYURBHAI KALUBHAI At: Sachana...") when name is already at top
        if name_upper and line.strip().upper().startswith(name_upper):
            rest = line.strip()[len(name_at_top.strip()):].strip()
            if re.match(r"^(At\s*:|Ta\s*:|Dist\s*:|\d)", rest, re.IGNORECASE):
                continue
        # Drop declaration and sign-off lines (unwanted at end of CV) — check anywhere in line (e.g. bullet text)
        if re.search(r"I\s+hereby\s+solemnly\s+declare|above\s+is\s+true\s+to\s+the\s+best\s+of\s+my\s+knowledge", line, re.I):
            continue
        if re.search(r"Your(s)?\s+Faithfully|Yours\s+Sincerely", line, re.I):
            continue
        if re.match(r"^\s*(Place\s*[:\-]|Date\s*[:\-])\s*", line, re.I):
            continue
        if re.match(r"^\s*Mr\.\s+[A-Za-z\s]+$", line) or re.search(r"Your\s+Faithfully\s+Mr\.", line, re.I):
            continue
        lines.append(line)
    return "\n".join(lines)


def _format_line_for_display(line: str) -> str:
    """Strip leading bullet chars and optional 'Label :' prefix for clean display."""
    line = line.strip()
    line = _strip_bad_chars(line)
    line = _BULLET_LINE.sub("", line).strip()
    # Remove common "Label : " prefix so we show value only (clearer CV)
    line = re.sub(r"^(Name|Email|Phone|Location|DOB|Address)\s*:\s*", "", line, flags=re.IGNORECASE)
    return line.strip()


def _is_bullet_line(line: str) -> bool:
    """True if line looks like a list item."""
    return bool(_BULLET_LINE.match(line.strip())) or line.strip().startswith(("-", "•", "*", "·", "–", "—"))


def _remove_duplicate_career_objective_from_text(text: str, career_objective: str) -> str:
    """Remove the career objective paragraph from text so it is not duplicated in Experience & Background."""
    if not text or not (career_objective or "").strip():
        return text
    normalized_obj = re.sub(r"\s+", " ", career_objective.strip())
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    out_lines = []
    i = 0
    while i < len(lines):
        line = lines[i]
        normalized_line = re.sub(r"\s+", " ", line)
        if normalized_line == normalized_obj:
            i += 1
            continue
        # Check if this line plus following lines form the career objective
        acc = normalized_line
        j = i + 1
        while j < len(lines) and acc != normalized_obj and len(acc) < len(normalized_obj) + 10:
            acc += " " + re.sub(r"\s+", " ", lines[j])
            j += 1
        if acc.strip() == normalized_obj:
            i = j
            continue
        out_lines.append(line)
        i += 1
    return "\n".join(out_lines)


# Section header pattern (for parsing)
_SECTION_HEADER = re.compile(
    r"^(CAREER\s*OBJECTIVE|OBJECTIVE|CAREER\s*GOAL|PERSONAL\s*DETAILS|"
    r"EXPERIENCE|WORK\s*EXPERIENCE|EDUCATION|SKILLS|QUALIFICATION|REFERENCES?)\s*[:\s]*$",
    re.IGNORECASE,
)
# Start of next section — stop Career Objective when we see these (line starts with)
_AFTER_OBJECTIVE_STOP = re.compile(
    r"^\s*(Personal\s*information|Personal\s*details|Name\s*:|Address\s*:|"
    r"Contact\s*no\s*:|Email\s*id\s*:|Date\s*of\s*birth|Educational\s*information|"
    r"Experience|Work\s*experience|Qualifications?|Skills\s*[:\s]|SR\s*NO\.)",
    re.IGNORECASE,
)
_LABEL_VALUE = re.compile(r"^(Name|DOB|Date\s*of\s*Birth|D\.?O\.?B\.?|Address|Email|Phone)\s*[:\-]\s*(.+)$", re.IGNORECASE)
_AT_COMPANY = re.compile(r"^(?:At|Company)\s*[:\-]\s*(.+)$", re.IGNORECASE)
_YEARS_EXPERIENCE = re.compile(r"(\d+)\s*(?:year|yr)s?\s*(?:experience|exp\.?)?", re.IGNORECASE)
_DATE_PATTERN = re.compile(r"\b(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})\b")


def parse_cv_into_structured_fields(text: str) -> Dict[str, Any]:
    """
    Parse raw extracted CV text into structured fields: name, dob, career_objective,
    experience (list of {title, company, location, start_date, end_date, description}).
    """
    result: Dict[str, Any] = {"career_objective": "", "dob": "", "experience": []}
    if not text or not text.strip():
        return result

    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    lines = [_strip_bad_chars(ln) for ln in lines]

    # Label : value extraction (Name, DOB, etc.)
    for line in lines:
        m = _LABEL_VALUE.match(line)
        if m:
            label, value = m.group(1).strip(), m.group(2).strip()
            value = _strip_bad_chars(value)
            if re.match(r"name", label, re.I):
                result["name"] = value
            elif re.match(r"dob|date\s*of\s*birth|d\.?o\.?b", label, re.I):
                result["dob"] = value

    # Career Objective: only the paragraph(s) after "Career Objective" / "Objective", stop at Personal info/Name/Address/Education etc.
    in_objective = False
    objective_lines = []
    for i, line in enumerate(lines):
        if _SECTION_HEADER.match(line):
            if re.match(r"^(CAREER\s*OBJECTIVE|OBJECTIVE|CAREER\s*GOAL)\s*", line, re.I):
                in_objective = True
                continue
            if in_objective:
                break
        if in_objective:
            # Stop as soon as we hit next section (e.g. "Personal information Name :" or "Name :" or "Address :")
            if _AFTER_OBJECTIVE_STOP.match(line):
                break
            objective_lines.append(line)
    if objective_lines:
        obj_text = " ".join(objective_lines)
        # If "Personal information" or "Name :" etc. appear in the same block, keep only the objective part
        stop_match = re.search(
            r"\s+(Personal\s*information|Personal\s*details|Name\s*:|Address\s*:|"
            r"Contact\s*no\s*:|Email\s*id\s*:|Date\s*of\s*birth|Educational\s*information)\s*",
            obj_text,
            re.IGNORECASE,
        )
        if stop_match:
            obj_text = obj_text[: stop_match.start()].strip()
        result["career_objective"] = obj_text

    # Experience block: after "Experience" / "Work Experience", collect lines and try to structure
    in_experience = False
    exp_lines = []
    for i, line in enumerate(lines):
        if _SECTION_HEADER.match(line) and re.match(r"^(EXPERIENCE|WORK\s*EXPERIENCE)\s*", line, re.I):
            in_experience = True
            continue
        if in_experience:
            if _SECTION_HEADER.match(line) and not re.match(r"^(EXPERIENCE|WORK)", line, re.I):
                break
            exp_lines.append(line)

    if exp_lines:
        # Build experience entries: detect company (At: X), dates, role, duties as bullets
        job: Dict[str, Any] = {"title": "", "company": "", "location": "", "start_date": "", "end_date": "", "description": []}
        for line in exp_lines:
            at_m = _AT_COMPANY.match(line)
            if at_m:
                if job.get("company") or job.get("title") or job.get("description"):
                    result["experience"].append(job)
                    job = {"title": "", "company": "", "location": "", "start_date": "", "end_date": "", "description": []}
                job["company"] = _strip_bad_chars(at_m.group(1))
                continue
            date_m = _DATE_PATTERN.search(line)
            if date_m:
                if not job.get("start_date"):
                    job["start_date"] = date_m.group(1)
            years_m = _YEARS_EXPERIENCE.search(line)
            if years_m:
                before = _strip_bad_chars(line[: years_m.start()].strip().rstrip(","))
                if before and not job.get("title"):
                    job["title"] = before
                elif not job.get("title"):
                    job["title"] = "Professional"
            # Location: "At: Sanand GIDC" or "Location: X" style
            if re.match(r"^(?:at|location)\s*[:\-]\s*", line, re.I):
                loc = re.sub(r"^(?:at|location)\s*[:\-]\s*", "", line, flags=re.I).strip()
                if loc and not _DATE_PATTERN.match(loc):
                    job["location"] = _strip_bad_chars(loc)
            # Duty lines: longer descriptive lines or bullet-like
            elif len(line) > 40 and "experience" not in line.lower() and not at_m:
                job["description"].append(_strip_bad_chars(line))
            elif 15 < len(line) <= 40 and not at_m and not date_m and "year" not in line.lower():
                if not job.get("title") and line[0].isupper():
                    job["title"] = _strip_bad_chars(line)
                elif line not in (job.get("description") or []):
                    job["description"].append(_strip_bad_chars(line))
        if job.get("company") or job.get("title") or job.get("description"):
            result["experience"].append(job)

    # Fallback: no "Experience" section but "Company Name :- X" appears (e.g. after Interested Field)
    if not result["experience"]:
        duty_pat = re.compile(r"^(?:I\s+am\s+|Performing|Operation|Operat|Fill\s+the|Perform\s+|Calibration|Manufacturing|Batch\s+|CIP|SIP|Operating\s+of)", re.I)
        for i, line in enumerate(lines):
            m = re.match(r"^(?:Company\s*Name|Company)\s*[:\-]+\s*(.+)$", line, re.IGNORECASE)
            if m:
                company_part = _strip_bad_chars(m.group(1))
                parts = re.split(r"\s+-\s*", company_part, maxsplit=1)
                job = {"title": "", "company": parts[0].strip() if parts else company_part, "location": parts[1].strip() if len(parts) > 1 else "", "start_date": "", "end_date": "", "description": []}
                for j in range(i + 1, len(lines)):
                    l = lines[j]
                    if re.search(r"I\s+hereby\s+solemnly\s+declare", l, re.I):
                        break
                    if re.match(r"^Interested\s*Field", l, re.I):
                        continue
                    date_m = _DATE_PATTERN.search(l)
                    if date_m and not job["start_date"]:
                        job["start_date"] = date_m.group(1)
                    years_m = _YEARS_EXPERIENCE.search(l)
                    if years_m and not job["title"]:
                        job["title"] = "Associate Technician" if re.search(r"technician|apprentice", l, re.I) else "Professional"
                    if duty_pat.match(l) or (len(l) > 35 and re.search(r"\b(operating|performing|operation|manufacturing|CIP|SIP|batch)\b", l, re.I)):
                        job["description"].append(_strip_bad_chars(l))
                    if re.match(r"^\s*At\s*[:\-]", l, re.I):
                        job["location"] = job["location"] or _strip_bad_chars(re.sub(r"^\s*At\s*[:\-]\s*", "", l, flags=re.I))
                if job.get("company") or job.get("description"):
                    result["experience"].append(job)
                break

    return result


# Patterns for reorganizing Experience & Background
_COMPANY_NAME_LINE = re.compile(r"^(?:Company\s*Name|Company)\s*[:\-]+\s*(.+)$", re.IGNORECASE)
_INTERESTED_FIELD_LINE = re.compile(r"^Interested\s*Field\s*[:\-]?\s*(.*)$", re.IGNORECASE)
_DUTY_LINE = re.compile(
    r"^(?:I\s+am\s+|I\s+have\s+|Performing|Operation|Operat|Fill\s+the|Filling|"
    r"Perform\s+|Calibration|Manufacturing|Batch\s+|CIP|SIP|Operating\s+of)",
    re.IGNORECASE,
)
_DECLARATION_START = re.compile(r"^I\s+hereby\s+solemnly\s+declare", re.IGNORECASE)


def _write_experience_background_content(document: Document, cleaned_text: str) -> None:
    """Write Experience & Background reorganized: Interested Field, Company/role, duties as bullets. Declaration/sign-off omitted."""
    lines = [ln.strip() for ln in cleaned_text.splitlines() if ln.strip()]
    i = 0
    while i < len(lines):
        line = lines[i]
        line_clean = _format_line_for_display(line)
        if not line_clean:
            i += 1
            continue
        # Declaration and sign-off: skip completely (including when they appear as bullets)
        if re.search(r"I\s+hereby\s+solemnly\s+declare", line_clean, re.I):
            j = i + 1
            while j < len(lines):
                next_clean = _format_line_for_display(lines[j])
                is_decl_or_signoff = next_clean and (
                    re.search(r"above\s+is\s+true\s+to\s+the\s+best|Your(s)?\s+Faithfully|Yours\s+Sincerely", next_clean, re.I)
                    or re.match(r"^\s*Mr\.\s+[A-Za-z\s]+$", next_clean)
                )
                if is_decl_or_signoff:
                    j += 1
                else:
                    break
            i = j
            continue
        if re.search(r"above\s+is\s+true\s+to\s+the\s+best\s+of\s+my\s+knowledge", line_clean, re.I):
            i += 1
            continue
        if re.search(r"Your(s)?\s+Faithfully|Yours\s+Sincerely", line_clean, re.I):
            i += 1
            continue
        if re.match(r"^\s*Mr\.\s+[A-Za-z\s]+$", line_clean):
            i += 1
            continue
        # Interested Field: subheading + value
        if _INTERESTED_FIELD_LINE.match(line):
            val = _INTERESTED_FIELD_LINE.sub(r"\1", line_clean).strip()
            if not val and i + 1 < len(lines):
                val = _format_line_for_display(lines[i + 1])
                i += 1
            if val:
                p = document.add_paragraph()
                run = p.add_run("Interested Field: " + val)
                _set_run_font(run, size_pt=BODY_PT, bold=True)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = LINE_SPACING
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            i += 1
            continue
        # Company Name :- X - Location
        if _COMPANY_NAME_LINE.match(line_clean):
            company_part = _COMPANY_NAME_LINE.sub(r"\1", line_clean).strip()
            p = document.add_paragraph()
            run = p.add_run(company_part)
            _set_run_font(run, size_pt=BODY_PT, bold=True)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = LINE_SPACING
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            i += 1
            continue
        # Duty/responsibility lines as bullets (I am operating, Performing, Operation of, etc.)
        if _DUTY_LINE.match(line_clean) or (len(line_clean) > 35 and re.search(r"\b(operating|performing|operation|manufacturing|CIP|SIP|batch)\b", line_clean, re.I)):
            _ats_bullet(document, line_clean)
            i += 1
            continue
        if _is_bullet_line(line):
            _ats_bullet(document, line_clean)
            i += 1
            continue
        # Experience intro (dates, duration, role): normal paragraph
        if re.search(r"\d+\s*(?:year|yr)s?\s*experience|\d{1,2}/\d{1,2}/\d{2,4}|associate|technician|apprentice", line_clean, re.I) and len(line_clean) < 120:
            _ats_body_paragraph(document, line_clean, space_after=4)
            i += 1
            continue
        # Short line + next line merge
        if len(line_clean) <= 50 and i + 1 < len(lines) and not _is_bullet_line(lines[i + 1]):
            next_clean = _format_line_for_display(lines[i + 1])
            if next_clean and len(next_clean) > 20:
                _ats_body_paragraph(document, line_clean + " " + next_clean, space_after=4)
                i += 2
                continue
        _ats_body_paragraph(document, line_clean, space_after=4)
        i += 1


def build_cv_doc(profile: Dict[str, Any]) -> Document:
    """Build an ATS-optimized, well-structured CV (single column, standard sections)."""
    document = Document()

    # Standard margins for ATS (1 inch)
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    name = profile.get("name", "")
    phone = profile.get("phone", "")
    email = profile.get("email", "")
    location = profile.get("location", "")
    dob = profile.get("dob", "")
    current_position = profile.get("current_position", "")
    current_company = profile.get("current_company", "")
    summary = profile.get("summary", "")
    career_objective = profile.get("career_objective", "")
    skills = profile.get("skills", [])
    experience = profile.get("experience", [])
    education = profile.get("education", [])
    original_text = profile.get("original_text", "")

    # ----- HEADER: Name -----
    if name:
        p = document.add_paragraph()
        run = p.add_run(name)
        _set_run_font(run, size_pt=NAME_PT, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # ----- Contact (single line, subtle font) -----
    contact_parts = []
    if phone:
        contact_parts.append(phone)
    if email:
        contact_parts.append(email)
    if location:
        contact_parts.append(location)
    if dob:
        contact_parts.append(f"DOB: {dob}")
    if contact_parts:
        p = document.add_paragraph()
        run = p.add_run("  ·  ".join(contact_parts))
        _set_run_font(run, size_pt=CONTACT_PT, bold=False)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # ----- Position at Company (no literal label) -----
    if current_position or current_company:
        current_parts = []
        if current_position:
            current_parts.append(current_position)
        if current_company:
            current_parts.append(f"at {current_company}")
        if current_parts:
            p = document.add_paragraph()
            run = p.add_run(" ".join(current_parts))
            _set_run_font(run, size_pt=CONTACT_PT, bold=False)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Divider line between header and body
    _add_horizontal_rule(document, space_after=14)

    # ----- CAREER OBJECTIVE -----
    if career_objective:
        _ats_section_heading(document, "Career Objective")
        obj_clean = _strip_bad_chars(career_objective).strip()
        if obj_clean:
            _ats_body_paragraph(document, obj_clean, space_after=8)

    # ----- PROFESSIONAL SUMMARY -----
    summary_text = summary
    if (current_position or current_company) and not summary_text:
        parts = []
        if current_position:
            parts.append(current_position)
        if current_company:
            parts.append(f"at {current_company}")
        summary_text = " ".join(parts) + "."
    if summary_text:
        _ats_section_heading(document, "Professional Summary")
        _ats_body_paragraph(document, summary_text, space_after=8)

    # ----- EXPERIENCE & BACKGROUND (cleaned, career objective removed to avoid duplication) -----
    if original_text:
        cleaned = _clean_extracted_cv_text(original_text, name_at_top=name)
        cleaned = _remove_duplicate_career_objective_from_text(cleaned, career_objective)
        if cleaned:
            _ats_section_heading(document, "Experience & Background")
            _write_experience_background_content(document, cleaned)
        original_text = ""

    # ----- WORK EXPERIENCE -----
    if experience:
        _ats_section_heading(document, "Work Experience")
        for job in experience:
            title = job.get("title", "")
            company = job.get("company", "")
            job_location = job.get("location", "")
            start_date = job.get("start_date", "")
            end_date = job.get("end_date", "")
            description_lines = job.get("description", [])

            # Job title and company on one line (ATS-friendly)
            line_parts = []
            if title:
                line_parts.append(title)
            if company:
                line_parts.append(f", {company}")
            if job_location:
                line_parts.append(f" — {job_location}")
            if line_parts:
                p = document.add_paragraph()
                run = p.add_run("".join(line_parts))
                _set_run_font(run, size_pt=BODY_PT, bold=True)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = LINE_SPACING
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

            if start_date or end_date:
                period = f"{start_date} – {end_date}".strip(" –")
                if period:
                    _ats_body_paragraph(document, period, space_after=4)

            for line in description_lines:
                if line.strip():
                    _ats_bullet(document, line.strip())
        # Extra space after last job
        document.add_paragraph()

    # ----- EDUCATION -----
    if education:
        _ats_section_heading(document, "Education")
        for edu in education:
            degree = edu.get("degree", "")
            institution = edu.get("institution", "")
            edu_location = edu.get("location", "")
            year = edu.get("year", "") or edu.get("end_date", "")

            line_parts = []
            if degree:
                line_parts.append(degree)
            if institution:
                line_parts.append(f", {institution}")
            if edu_location:
                line_parts.append(f" — {edu_location}")
            if line_parts:
                p = document.add_paragraph()
                run = p.add_run("".join(line_parts))
                _set_run_font(run, size_pt=BODY_PT, bold=True)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = LINE_SPACING
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            if year:
                _ats_body_paragraph(document, year, space_after=6)
        document.add_paragraph()

    # ----- SKILLS -----
    if skills:
        _ats_section_heading(document, "Skills")
        # ATS-friendly: comma-separated or one bullet per line (we use bullets for clarity)
        for skill in skills:
            if isinstance(skill, str) and skill.strip():
                _ats_bullet(document, skill.strip())
        document.add_paragraph()

    return document


def save_docx_and_optional_pdf(doc: Document, output_docx: Path, generate_pdf: bool) -> None:
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))

    if generate_pdf:
        if docx2pdf_convert is None:
            raise RuntimeError("docx2pdf is not installed; cannot generate PDF.")
        output_pdf = output_docx.with_suffix(".pdf")
        docx2pdf_convert(str(output_docx), str(output_pdf))


def main() -> None:
    parser = argparse.ArgumentParser(
        description=(
            "Scan an existing CV (PDF/image), load the latest profile data, "
            "and generate an updated CV in DOCX/PDF using a clean template."
        )
    )
    parser.add_argument(
        "--input",
        type=str,
        required=True,
        help="Path to the existing CV file (PDF or image).",
    )
    parser.add_argument(
        "--profile",
        type=str,
        required=True,
        help="Path to the profile JSON with the latest details.",
    )
    parser.add_argument(
        "--output-docx",
        type=str,
        required=True,
        help="Path for the generated DOCX file.",
    )
    parser.add_argument(
        "--output-pdf",
        action="store_true",
        help="If set, also generate a PDF version next to the DOCX.",
    )

    args = parser.parse_args()

    input_path = Path(args.input)
    profile_path = Path(args.profile)
    output_docx_path = Path(args.output_docx)

    # Extract text from input for inspection/logging (future use for smart diff)
    extracted_text = ""
    if input_path.suffix.lower() == ".pdf":
        try:
            extracted_text = extract_text_from_pdf(input_path)
        except Exception as e:
            print(f"Warning: Failed to extract text from PDF: {e}")
    else:
        # Placeholder for image OCR or other formats
        print("Non-PDF input detected; OCR-based extraction can be added here.")

    if extracted_text:
        print("\n--- Extracted text (truncated) ---\n")
        print(extracted_text[:2000])
        print("\n--- End of extracted text preview ---\n")

    profile = load_profile(profile_path)

    # Build DOCX from profile (source of truth for current position, phone, etc.)
    document = build_cv_doc(profile)
    save_docx_and_optional_pdf(document, output_docx_path, args.output_pdf)

    print(f"Generated DOCX: {output_docx_path}")
    if args.output_pdf:
        print(f"Generated PDF: {output_docx_path.with_suffix('.pdf')}")


if __name__ == "__main__":
    main()

